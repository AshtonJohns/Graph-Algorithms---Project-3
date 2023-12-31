import matplotlib.pyplot as plt
from matplotlib.pyplot import figure
import networkx as nx
import numpy as np
import random
from docx import Document
from docx.shared import Inches
import os

## produce graph without variation of node positions (when position is NOT defined)
# seed = 0
# random.seed(seed)
# np.random.seed(seed)
def addToDoc(document,task="",text="",tmpFile="",addparagraph=False,addFigure=False):
    if addparagraph: #adding bolded paragraph
        p = document.add_paragraph()
        p.add_run(text).bold = True
    if addFigure:
        document.add_heading(task,level=1)
        document.add_picture(tmpFile)

def generateGraph(G,pos,tmpFile,document,task,answer='',mainIllustration=False,wouldLikeToViewGraphs=True,node_size=1000,x=3,y=3): #genereate all kinds of graphs
    if mainIllustration:
        if task.startswith("task 2"): #directed
            nx.draw(G,pos=pos,with_labels=True,node_color="red",font_color="white",font_size=15,font_family="Times New Roman", font_weight="bold",width=3,edge_color="black")
        elif task.startswith("task 3"): #undirected, weighted
            nx.draw_networkx(G, pos)
            nx.draw_networkx_edge_labels(G, pos, edge_labels=nx.get_edge_attributes(G, 'weight')) 
        else: #undirected
            nx.draw_networkx(G,pos=pos,with_labels=True)
        plt.margins(0.2)
        #add image to docx file
        plt.savefig(tmpFile)
        # add figure to document
        if answer != '':
            task = answer
        addToDoc(document,task=task,tmpFile=tmpFile,addFigure=True)
    elif not mainIllustration and task.startswith("task 2"): #task 2 graphs
        if answer.startswith("DAG: linearized in topological order"):
            nx.draw(G,pos=pos,node_size=700,with_labels=True,width=3,connectionstyle='arc3, rad=0.5') #connectionstyle allows for the directed edges to be curved for viewing purposes
        else:
            nx.draw_networkx(G,node_size=node_size,width=3)
        plt.margins(0.2)
        plt.savefig(tmpFile)
        if answer != '':
            task = answer
        addToDoc(document,task=task,tmpFile=tmpFile,addFigure=True)
    elif not mainIllustration and task.startswith("task 3"): #task 3 graphs
        if answer.startswith("Kruskal's algorithm: mst graph"):
            edges = G.edges()
            colors = [G[u][v]['color'] for u,v in edges]
            nx.draw(G, pos, edge_color=colors, width=2,with_labels=True)
            nx.draw_networkx_edge_labels(G,pos=pos,edge_labels=nx.get_edge_attributes(G, 'weight'))
        elif answer.startswith("Dijkstra’s algorithm: spt graph"):
            edges = G.edges()
            colors = [G[u][v]['color'] for u,v in edges] #get the color for each edge
            #weights = [G[u][v]['weight'] for u,v in edges]
            nx.draw(G, pos, edge_color=colors, width=2,with_labels=True) #assign the correct colors to each of the edges (regular paths will be black, spt paths will be red)
            nx.draw_networkx_edge_labels(G,pos=pos,edge_labels=nx.get_edge_attributes(G, 'weight')) # assign the weights to each edge
        plt.margins(0.2)
        plt.savefig(tmpFile)
        if answer != '':
            task = answer
        addToDoc(document,task=task,tmpFile=tmpFile,addFigure=True)
    else: # compare 
        nx.draw(G,pos=pos,with_labels=True,node_color="green",node_size=node_size,font_color="white",font_size=10,font_family="Times New Roman", font_weight="bold",width=2,edge_color="black")
        plt.margins(0.2)
        #save to figure
        figure = plt.gcf()
        figure.set_size_inches(x,y)
        #add image to docx file
        plt.savefig(tmpFile)
        # add figure to document
        if answer != '':
            task = answer
        addToDoc(document,task=task,tmpFile=tmpFile,addFigure=True)

    if wouldLikeToViewGraphs: # if the user wants to compare the figures
        plt.show(block=False) # display graph
        # inform user to press any key to move on
        #print(pressAnyKey)
        input("\n---Press Enter to continue---")
        plt.clf()
        plt.close()
    else:
        plt.clf()
        plt.close()

#-----------REVISE : was used to draw graphs side by side, but this was determined to not be a valid approach for this task.-------------
def generateGraphSideBySideForComparison(document,G1,G2,pos,task,tmpFile,operation='',answer='',operation2='',wouldLikeToViewGraphs=True):
    image1 = ""
    image2 = ""
    if operation.startswith("compare spt and mst"):
        image1 = "spt" + tmpFile
        edges = G1.edges()
        colors = [G1[u][v]['color'] for u,v in edges]
        nx.draw(G1, pos, edge_color=colors, width=2,with_labels=True)
        nx.draw_networkx_edge_labels(G1,pos=pos,edge_labels=nx.get_edge_attributes(G1, 'weight'))
        plt.savefig(image1)
    
        image2 = "mst" + tmpFile
        edges = G2.edges()
        colors = [G2[u][v]['color'] for u,v in edges] #get the color for each edge
        #weights = [G[u][v]['weight'] for u,v in edges]
        nx.draw(G2, pos, edge_color=colors, width=2,with_labels=True) #assign the correct colors to each of the edges (regular paths will be black, spt paths will be red)
        nx.draw_networkx_edge_labels(G2,pos=pos,edge_labels=nx.get_edge_attributes(G2, 'weight')) # assign the weights to each edge
        plt.savefig(image2)

    if wouldLikeToViewGraphs:
        # Load the images
        plt.clf()
        plt.close()
        image1 = plt.imread(image1)
        image2 = plt.imread(image2)
        # Create a figure and axes
        fig, axes = plt.subplots(1, 2)
        #plt.subplots_adjust(left=0.1, bottom=0.1, right=0.9, top=0.9)
        # Plot the images
        axes[0].imshow(image1)
        axes[1].imshow(image2)
        tmpFile = "sideByside.png"
        plt.savefig(tmpFile)
        # add figure to document
        if answer != '':
            task = answer
        addToDoc(document,task=task,tmpFile=tmpFile,addFigure=True)
        # Show the figure
        plt.show(block=False) # display graph
        # inform user to press any key to move on
        #print(pressAnyKey)
        input("\n---Press Enter to continue---")
        plt.clf()
        plt.close()
    


def viewGraphs():
    cont = True
    viewFile = False
    answerQuestion = False
    while cont:
        print("\n\nWould you like to answer this question? \n1. Yes\n2. No\n")
        answerQChoice = input("Choice: ")
        if answerQChoice.__eq__("1"):
            answerQuestion = True
            cont = False
            # print("\nWould you like to see the graphs for comparison? (If not, SAVE THE FILE to view later)\n1. Yes\n2. No\n")
            # choice = input("Choice: ")
            # if choice.__eq__("1"):
            #     viewFile = True
            #     cont = False
            # elif choice.__eq__("2"):
            #     print("\n\nOkay, you will not see the comparisons\n\n")
            #     cont = False
        elif answerQChoice.__eq__("2"):
            cont = False
        else:
            print("Please select a valid option")
    return viewFile,answerQuestion

def getConnectedComponents(G):
    conn_comp_gen = nx.connected_components(G) #gets connected components from graph
    #list comprehension 
    temp_list = list(conn_comp_gen)
    conn_comp = []
    for connected_components_list in temp_list:
        conn_comp.append(list(connected_components_list))
    return conn_comp #a list of each of the connected components 

def getStronglyConnectedComponents(G):
    scc_gen = nx.strongly_connected_components(G) #gets scc 
    #list comprehension
    temp_list = list(scc_gen)
    scc = []
    for scc_list in temp_list:
        scc.append(list(scc_list))
    return scc #list of all scc (eg, [1],[5,6,7])

def createDigraphWithSCC(G):
    testter = nx.condensation(G) #get the condensation of G is the graph with each of the scc contracted into a single node
    node_data = testter.nodes.data()
    map = {}
    for node_data_stuff in node_data: #turns generator into a hashmap, with keys being nodes (representations of scc), and values being scc
        test_str = str(node_data_stuff[1]['members'])
        test_str = test_str[1:]
        test_str = test_str[:-1]
        map[node_data_stuff[0]] = test_str
    testter = nx.relabel_nodes(testter,map) #relabels the node representations of scc (eg 1 -> [1,2,3], 2 -> [4])
    return testter #digraph of scc

def createDAG(G):
    testter = nx.dag_to_branching(G) # get creates DAG
    node_data = testter.nodes.data()
    map = {}
    for node_data_stuff in node_data: #since this DAG has nodes not the same as digraph, get the mapping
        test_str = str(node_data_stuff[1]['source'])
        map[node_data_stuff[0]] = test_str
    testter = nx.relabel_nodes(testter,map) #relabels nodes to accurately depict the nodes from digraph

    node_data = testter.nodes.data()
    pos = {}
    i = 0
    for node_data_stuff in node_data: # linearize the DAG in the correct order -> builds a hashmap with positions horizontal incrementing by 2.5 places
        pos[node_data_stuff[0]] = (i,5)
        i += 2.5
    return testter,pos #returns DAG and positions for the nodes

def generateGraphForSPT(document,G_weighted_undirected,pos,task,answer,tmpFile,start_node,viewGraph=True,affectDocument=True,target_node=''):
    path_edge_color = 'red'
    nodes = G_weighted_undirected.nodes()

    if task.__eq__("task 3 question 1") or task.__eq__("task 3 question 3"):
        spt = nx.single_source_dijkstra_path(G_weighted_undirected,start_node,weight='weight')

        for key,value in spt.items():
            message1 = "Dijkstra's algorithm found the shortest path from " + start_node + " to " + key + ":\n"
            message1 += str(value) + "\n"
            message1 += "\nShortest path tree: \n"
            G_s = G_weighted_undirected.copy()
            
            if len(value) == 2:
                edges_data = G_s.get_edge_data(value[0],value[1])
                G_s.add_edge(value[0],value[1],color=path_edge_color,weight=edges_data.get('weight'))

            elif len(value) > 2:
                for x in range(len(value)):
                    if not value[-2].__eq__(value[-3]):
                        index = value.index(value[x+x+1])
                        value.insert(index+1,value[x+x+1])
                #convert to 2d
                edges_2d_list = np.array(value).reshape((len(value))//2,2)
                for x in range(len(edges_2d_list)):
                    edges_data = G_s.get_edge_data(edges_2d_list[x][0],edges_2d_list[x][1])
                    G_s.add_edge(edges_2d_list[x][0],edges_2d_list[x][1],color=path_edge_color,weight=edges_data.get('weight'))
            # add question to document and generate graph
            if affectDocument:
                print(message1)
                addToDoc(document,text=message1,addparagraph=True)
                generateGraph(G_s,pos,tmpFile,document,task,answer=answer,mainIllustration=False,wouldLikeToViewGraphs=viewGraph)
                document.add_page_break()
    elif task.__eq__("task 3 question 4"):
        spt = nx.dijkstra_path(G_weighted_undirected,start_node,target=target_node,weight='weight')

        value = spt

        message1 = "Dijkstra's algorithm found the shortest path from " + start_node + " to " + target_node + ":\n"
        message1 += str(value) + "\n"
        message1 += "\nShortest path tree: \n"
        G_s = G_weighted_undirected.copy()
        
        if len(value) == 2:
            edges_data = G_s.get_edge_data(value[0],value[1])
            G_s.add_edge(value[0],value[1],color=path_edge_color,weight=edges_data.get('weight'))

        elif len(value) > 2:
            for x in range(len(value)):
                if not value[-2].__eq__(value[-3]):
                    index = value.index(value[x+x+1])
                    value.insert(index+1,value[x+x+1])
            #convert to 2d
            edges_2d_list = np.array(value).reshape((len(value))//2,2)
            for x in range(len(edges_2d_list)):
                edges_data = G_s.get_edge_data(edges_2d_list[x][0],edges_2d_list[x][1])
                G_s.add_edge(edges_2d_list[x][0],edges_2d_list[x][1],color=path_edge_color,weight=edges_data.get('weight'))
        # add question to document and generate graph
        if affectDocument:
            print(message1)
            addToDoc(document,text=message1,addparagraph=True)
            generateGraph(G_s,pos,tmpFile,document,task,answer=answer,mainIllustration=False,wouldLikeToViewGraphs=viewGraph)
    
    return G_s #graph object

def generateGraphForMST(document,G_weighted_undirected,pos,task,answer,tmpFile,viewGraph=True,affectDocument=True):
    mst = nx.minimum_spanning_tree(G_weighted_undirected,weight='weight',algorithm='kruskal')
    all_edges = mst.edges()
    path_edge_color = 'red'

    edges_2d_list = []
    for x in all_edges:
        edges_2d_list.append([x[0],x[1]])

    G_s = G_weighted_undirected.copy()

    for x in range(len(edges_2d_list)):
        edges_data = G_s.get_edge_data(edges_2d_list[x][0],edges_2d_list[x][1])
        G_s.add_edge(edges_2d_list[x][0],edges_2d_list[x][1],color=path_edge_color,weight=edges_data.get('weight'))
        message = "\nTo get to node " + str(edges_2d_list[x][1]) + ", the closest node with least weight is " + str(edges_2d_list[x][0])
        message += "\nGraphically represented:\n"
        if x == len(edges_2d_list) - 1:
            message += "\n!---ALL NODES AND PATHS HAVE BEEN MADE WITH THE LEAST WEIGHT USING KRUSKAL'S ALGORITHM---!:\n"
        if affectDocument:
            print(message)
            addToDoc(document,text=message,addparagraph=True)
            generateGraph(G_s,pos,tmpFile,document,task,answer=answer,mainIllustration=False,wouldLikeToViewGraphs=viewGraph)
            document.add_page_break()
    
    return G_s #graph object

def task1(document):
    #current task
    task = "Task 1"
    #temp file for graphs
    tmpFile = "graph.png"
    # move on message
    #pressAnyKey = "\n\n" + task + "graph\n\n" + "\n---PRESS ANY KEY TO MOVE ON!---\n\n"
    # questions
    question1 = "\n\nQUESTION 1:\nStarting from any vertex, can DFS and BFS find all connected components of an undirected graph?\n\n"
    question2 = "\n\nQUESTION 2:\nCan both BFS and DFS determine if there is a path between two given nodes?\n\n"
    question3 = "\n\nQUESTION 3:\nProvided that there is a path between two vertices u and v in the graph. If started from u, do DFS and BFS always find exactly the same path to v?\n\n"
    #create graph
    G=nx.Graph()

    G_undirected_edges = [("A", "B"),
                          ("A", "F"),
                          ("A", "E"),
                        ("B", "F"),
                        ("B", "C"),
                        ("C", "D"),
                        ("C", "G"),
                        ("D", "G"),
                        ("G", "J"),
                        ("E", "F"),
                        ("E", "I"),
                        ("F", "I"),
                        ("I", "J"),
                        ("I", "M"),
                        ("M", "N"),
                        ("H", "K"),
                        ("K", "O"),
                        ("K", "L"),
                        ("L", "H"),
                        ("L", "P")]
    #add edges
    G.add_edges_from(G_undirected_edges)

    G_undirected = G.to_undirected()

    #get connected components
    conn_comp = getConnectedComponents(G)
    
    # list_connected_components_1 = ['A','B','C','D','E','F','G','I','J','M','N']
    # list_connected_components_2 = ['H','K','L','O','P']

    pos={ # positions for nodes
        
        "A":(0,10),
        "B":(2.5,10),
        "C":(5,10),
        "D":(7.5,10),
        "E":(0,7.5),
        "F":(2.5,7.5),
        "G":(5,7.5),
        "H":(7.5,7.5),
        "I":(0,5),
        "J":(2.5,5),
        "K":(5,5),
        "L":(7.5,5),
        "M":(0,2.5),
        "N":(2.5,2.5),
        "O":(5,2.5),
        "P":(7.5,2.5),
    }

    #
    # ######################### --------------- QUESTION 1 ---------------#########################  #
    # 
    print(question1)
    task = "task 1 question 1"

    # can DFS & BFS find all components of undirected graph?

    #ask if the user wants to compare the graphs, or just view it later in the file
    view,answerQuestion = viewGraphs()     

    if answerQuestion:
        # generate undirected graph 
        generateGraph(G_undirected,pos,tmpFile,document,task,mainIllustration=True,wouldLikeToViewGraphs=True)
        # add question to document
        addToDoc(document,text=question1,addparagraph=True)
        document.add_page_break()
        
        for node in G_undirected.nodes():#loop over nodes
            #list for storing results for DFS & BFS
            dfs = []
            bfs = []
            for current_node in nx.dfs_tree(G_undirected,node):
                dfs.append(current_node)
            for current_node in nx.bfs_tree(G_undirected,node):
                bfs.append(current_node)

            #determine which connected components to look at 
            current_conn_comp = []
            for known_conn_comp_part in conn_comp:
                if node in known_conn_comp_part:
                    current_conn_comp = known_conn_comp_part
            # message1
            message1 = "\n\nDFS\nStart node: " + node + "\nConnected Components found by DFS from the start node: \n" + str(dfs) + "\nKnown Connected Components: \n" 
            for each_part in conn_comp:
                message1 += str(each_part) + "\n"           
            message1 += "Known Connected Components in Question: \n" + str(current_conn_comp) \
                        + "\nDid DFS find all connected components from the start node?\n" +"(DFS)" + str(dfs) + "\n=?\n" + str(current_conn_comp) + "\nResult: "
            result = set(current_conn_comp).__eq__(set(dfs))
            message1 += str(result)
            print(message1)
            addToDoc(document,text=message1,addparagraph=True) 
            # message2
            message2 = "\n\nBFS\nStart node: " + node + "\nConnected Components found by BFS from the start node: \n" + str(bfs) + "\nKnown Connected Components: \n" \
                        + str(conn_comp) + "\nKnown Connected Components in Question: \n" + str(current_conn_comp) \
                        + "\nDid BFS find all connected components from the start node?\n" +"(BFS)" + str(dfs) + "\n=?\n" + str(current_conn_comp) + "\nResult: "
            result = set(current_conn_comp).__eq__(set(dfs))
            message2 += str(result)
            print(message2)
            addToDoc(document,text=message2,addparagraph=True)
            document.add_page_break()


            #REVISE : not used!
            # #list for storing results for DFS & BFS
            # dfs = []
            # bfs = []
            # #generate temp graph
            # temp_G_dfs = nx.Graph()
            # temp_G_bfs = nx.Graph()
            # previous_node = '' #stores previous node while doing dfs or bfs
            # for current_node in nx.dfs_tree(G_undirected,node):
            #     dfs.append(current_node)
            #     if previous_node != current_node and previous_node != '': #create the temp graph
            #         temp_G_dfs.add_edge(previous_node,current_node)
            #     # assign the previous node
            #     previous_node = current_node
            # previous_node = '' #reset for bfs
            # for current_node in nx.bfs_tree(G_undirected,node):
            #     bfs.append(current_node)
            #     if previous_node != current_node and previous_node != '': #create the temp graph for bfs
            #         temp_G_bfs.add_edge(previous_node,current_node)
            #     previous_node = current_node
            # #display both dfs & bfs figures for the user to see if they chose to view
            # generateGraphSideBySideForComparison(temp_G_dfs,temp_G_bfs,pos,tmpFile,document,task,pressAnyKey,wouldLikeToViewGraphs=view)
            # # determine which list of connected components the node belongs within
            # if node in list_connected_components_1:
            #     list_connected_components = list_connected_components_1
            # else:
            #     list_connected_components = list_connected_components_2
            # # print message / doc message
            # message = "\nKnown connected components: " + str(list_connected_components) + \
            #         "\nStarting from node: " + node + "..." + \
            #         "\nConnected components determined with DFS: " + str(dfs) + \
            #         "\nConnected components determined with BFS: " + str(bfs) + \
            #         "\nDid DFS and BFS result in getting all the connected components?" + \
            #         "\n(DFS)" + str(dfs) + "=?" + str(list_connected_components) + ": "
            # # compare with list of connected components
            # dfs.sort()
            # result = list_connected_components.__eq__(dfs)
            # message += str(result)
            # message += "\n(BFS)" + str(bfs) + "=?" + str(list_connected_components) + ": "
            # bfs.sort()
            # result = list_connected_components.__eq__(bfs)
            # message += str(result)
            # print(message)
            # addToDoc(document,text=message,addparagraph=True)
    
    #
    # ######################### --------------- QUESTION 2 ---------------#########################  #
    # 
    print(question2)
    # message = "\nTo answer this question, we will loop over all nodes, and find all possible paths, using DFS and BFS\n\n"
    # print(message)
    # addToDoc(document,text=message,addparagraph=True)
    # generate undirected graph 
    task = "task 1 question 2"
    # Can both BFS and DFS determine if there is a path between two given nodes?
    #ask if the user wants to compare the graphs, or just view it later in the file
    view,answerQuestion = viewGraphs() 

    if answerQuestion:
        generateGraph(G_undirected,pos,tmpFile,document,task,mainIllustration=True,wouldLikeToViewGraphs=True)
        # add question to document
        addToDoc(document,text=question2,addparagraph=True)
        document.add_page_break()

        for node in G_undirected.nodes(): #loop over nodes
            # message1
            message1 = "\n\nDFS\nStart node: " + node
            print(message1)
            addToDoc(document,text=message1,addparagraph=True)
            
            for current_node in nx.dfs_tree(G_undirected, node): #create a dfs_tree starting from node (eg, "A"). 
                if current_node != node: # A == A, B == B, etc (starting nodes)
                    # message2
                    message2 = "\nDFS found a path from start node " + node + " and end node " + current_node  #states the path found from the start node to the end node (current_node -> all nodes in dfs_tree)
                    print(message2)
                    addToDoc(document,text=message2,addparagraph=True)
            # message3
            message3 = "\n\nBFS\nStart node: " + node
            print(message3)
            addToDoc(document,text=message3,addparagraph=True)
            for current_node in nx.bfs_tree(G_undirected, node):
                if current_node != node:
                    # message4
                    message4 = "\nBFS found a path from start node " + node + " and end node " + current_node  
                    print(message4)
                    addToDoc(document,text=message4,addparagraph=True)
            document.add_page_break()
            
    #REVISE : not used!
    # if answerQuestion:
    #     for node in G_undirected.nodes():
    #         dfs_for_undirected = []
    #         # message1
    #         message1 = "\n\nDFS\nStart node: " + node
    #         print(message1)
    #         addToDoc(document,text=message1,addparagraph=True)
    #         for current_node in nx.dfs_tree(G_undirected,node):
    #             test_list.append(current_node)
    #             dfs_for_undirected.append(current_node)
    #             if current_node != node: # A == A, B == B, etc (starting nodes)
    #                 dfs_for_undirected.append(current_node)
    #         dfs_for_undirected.pop()
    #         dfs_for_undirected_to_path = []
    #         for x in range(len(dfs_for_undirected)-1):
    #             temp_list = []
    #             if x % 2 == 0:
    #                 temp_list.append([dfs_for_undirected[x],dfs_for_undirected[x+1]])
    #                 path = nx.Graph()
    #                 path.add_edges_from(temp_list)

    #                 dfs_for_undirected_to_path.append([dfs_for_undirected[x],dfs_for_undirected[x+1]])
    #                 #convert to 2d
    #                 #dfs_to_2d = np.array(dfs_for_undirected_to_path).reshape((len(dfs_for_undirected_to_path))//2,2)
    #                 temp_G_dfs = nx.Graph()
    #                 temp_G_dfs.add_edges_from(dfs_for_undirected_to_path)

    #                 generateGraph(path,pos,tmpFile,document,task,pressAnyKey,wouldLikeToViewGraphs=view,x=2,y=2,node_size=300)
    #                 generateGraph(temp_G_dfs,pos,tmpFile,document,task,pressAnyKey,wouldLikeToViewGraphs=view,x=2,y=2,node_size=500)
    #                 # message2
    #                 message2 = "\nEnd node: " + dfs_for_undirected[x+1] + "\nDFS found a path from start node " + node + " and end node " + dfs_for_undirected[x+1] + "\nNew Edge: " + str(temp_list) + "\nTotal path: \n" + str(dfs_for_undirected_to_path) 
    #                 print(message2)
    #                 addToDoc(document,text=message2,addparagraph=True)
    #                 document.add_page_break() 

    #
    # ######################### --------------- QUESTION 3 ---------------#########################  #
    # 
    print(question3)
    # generate undirected graph 
    task = "task 1 question 3"
    # Provided that there is a path between two vertices u and v in the graph. If started from u, do DFS and BFS always find exactly the same path to v?
    #ask if the user wants to compare the graphs, or just view it later in the file
    view,answerQuestion = viewGraphs() 

    if answerQuestion:
        generateGraph(G_undirected,pos,tmpFile,document,task,mainIllustration=True,wouldLikeToViewGraphs=True)
        # add question to document
        addToDoc(document,text=question3,addparagraph=True)
        document.add_page_break()

        #find the known connected component parts
        conn_comp_1 = conn_comp[0] #all connected nodes from spanning tree (conn_comp is a forest)
        conn_comp_2 = conn_comp[1] #all connected nodes from spanning tree (conn_comp is a forest)

        for node_A in conn_comp_1: #loop over all connected components already found 
            for node_B in conn_comp_1: #loop over all connected components already found
                if node_A != node_B: # have start node compare to all connected components (A -> B, A -> C, etc)
                    dfs = []
                    bfs = []
                    #message 1 
                    message1 = "\nKnown Connected Components: \n" + str(conn_comp_1) \
                             + "\nStart node: \n" + node_A + "\nNode to compare DFS and BFS path: \n" + node_B
                    for i in nx.dfs_tree(G_undirected, node_A): #make list as before for all connected components found by dfs
                        dfs.append(i)
                    message1 += "\nConnected Components found by DFS from the start node: \n" + str(dfs)
                    for i in nx.bfs_tree(G_undirected, node_A): #make list as before for all connected components found by bfs
                        bfs.append(i)
                    message1 += "\nConnected Components found by BFS from the start node: \n" + str(bfs)
                    dfs_last_node = dfs.index(node_B) + 1 #index of last node... (the +1 allows for a correct list slice)
                    bfs_last_node = bfs.index(node_B) + 1
                    dfs_path = dfs[:dfs_last_node]
                    message1 += "\n(DFS)Connected Components from start node to comparable node: \n" + str(dfs_path)
                    bfs_path = bfs[:bfs_last_node]
                    message1 += "\n(BFS)Connected Components from start node to comparable node: \n" + str(bfs_path) \
                            + "\nAre the paths exactly the same? \n" + "(DFS)" + str(dfs_path) + "\n=?\n" + "(BFS)" + str(bfs_path) \
                            + "\nResult: "
                    if dfs_path == bfs_path:
                        message1 += "True"
                    else: 
                        message1 += "False" 
                    print(message1)
                    addToDoc(document,text=message1,addparagraph=True)
                    document.add_page_break()
        # same as before, but with the second list of connected components 
        for node_A in conn_comp_2: 
            for node_B in conn_comp_2:
                if node_A != node_B:
                    dfs = []
                    bfs = []
                    #message 2 
                    message2 = "\nKnown Connected Components: \n" + str(conn_comp_2) \
                             + "\nStart node: \n" + node_A + "\nNode to compare DFS and BFS path: \n" + node_B
                    for i in nx.dfs_tree(G_undirected, node_A):
                        dfs.append(i)
                    message2 += "\nConnected Components found by DFS from the start node: \n" + str(dfs)
                    for i in nx.bfs_tree(G_undirected, node_A):
                        bfs.append(i)
                    message2 += "\nConnected Components found by BFS from the start node: \n" + str(bfs)
                    dfs_last_node = dfs.index(node_B) + 1
                    bfs_last_node = bfs.index(node_B) + 1
                    dfs_path = dfs[:dfs_last_node]
                    message2 += "\n(DFS)Connected Components from start node to comparable node: \n" + str(dfs_path)
                    bfs_path = bfs[:bfs_last_node]
                    message2 += "\n(BFS)Connected Components from start node to comparable node: \n" + str(bfs_path) \
                            + "\nAre the paths exactly the same? \n" + "(DFS)" + str(dfs_path) + "\n=?\n" + "(BFS)" + str(bfs_path) \
                            + "\nResult: "
                    if dfs_path == bfs_path:
                        message2 += "True"
                    else: 
                        message2 += "False" 
                    print(message2)
                    addToDoc(document,text=message2,addparagraph=True)
                    document.add_page_break()     

    #REVISE : not used!
    # if answerQuestion:
    #     for node in G_undirected.nodes():
    #         #list for storing results for DFS & BFS
    #         dfs = [] #stores all edges for DFS
    #         bfs = [] #stores all edges for BFS
    #         temp_list = [] #stores an edge
    #         # message1
    #         message1 = "\n\nDFS\nStart node: " + node
    #         print(message1)
    #         addToDoc(document,text=message1,addparagraph=True)
    #         for current_node in nx.dfs_tree(G_undirected,node):
    #             #generate graph containing all discovered edges
    #             temp_G_dfs = nx.Graph()
    #             #generate graph containing newly discovered edge
    #             path = nx.Graph()
    #             #add add nodes to create a new edge
    #             temp_list.append(current_node) 
    #             if len(temp_list) == 2: #once two nodes are connected...
    #                 dfs.extend(temp_list)
    #                 # add path from previous edge end node to start node of newly discovered edge
    #                 if len(dfs) > 2:
    #                     dfs.insert(dfs.index(dfs[-2]),dfs[-3])
    #                     dfs.insert(dfs.index(dfs[-2]),dfs[-2])
    #                 #convert to 2d
    #                 dfs_to_2d = np.array(dfs).reshape((len(dfs))//2,2) 
    #                 # create the total path for DFS
    #                 temp_G_dfs.add_edges_from(dfs_to_2d)
    #                 # create the newly discovered edge 
    #                 path.add_edge(temp_list[0],temp_list[1])
    #                 generateGraph(path,pos,tmpFile,document,task,pressAnyKey,wouldLikeToViewGraphs=view,x=2,y=2,node_size=500)
    #                 generateGraph(temp_G_dfs,pos,tmpFile,document,task,pressAnyKey,wouldLikeToViewGraphs=view,x=2,y=2,node_size=500)
    #                 # message2
    #                 message2 = "\nEnd node: " + current_node + "\nDFS found a path from start node " + node + " and end node " + current_node + "\nNew Edge: " + str(temp_list) + "\nTotal path: \n" + str(dfs_to_2d.flatten()) 
    #                 print(message2)
    #                 addToDoc(document,text=message2,addparagraph=True)
    #                 document.add_page_break()
    #                 temp_list = [] #clean edge
    #         # message3
    #         message3 = "\n\nBFS\nStart node: " + node
    #         print(message3)
    #         addToDoc(document,text=message3,addparagraph=True)
    #         for current_node in nx.bfs_tree(G_undirected,node):
    #             #generate graph containing all discovered edges
    #             temp_G_bfs = nx.Graph()
    #             #generate graph containing newly discovered edge
    #             path = nx.Graph()
    #             #add add nodes to create a new edge
    #             temp_list.append(current_node) 
    #             if len(temp_list) == 2: #once two nodes are connected...
    #                 bfs.extend(temp_list)
    #                 # add path from previous edge end node to start node of newly discovered edge
    #                 if len(bfs) > 2:
    #                     bfs.insert(bfs.index(bfs[-2]),bfs[-3])
    #                     bfs.insert(bfs.index(bfs[-2]),bfs[-2])
    #                 #convert to 2d
    #                 bfs_to_2d = np.array(bfs).reshape((len(bfs))//2,2) 
    #                 # create the total path for bfs
    #                 temp_G_bfs.add_edges_from(bfs_to_2d)
    #                 # create the newly discovered edge 
    #                 path.add_edge(temp_list[0],temp_list[1])
    #                 generateGraph(path,pos,tmpFile,document,task,pressAnyKey,wouldLikeToViewGraphs=view,x=2,y=2,node_size=500)
    #                 generateGraph(temp_G_bfs,pos,tmpFile,document,task,pressAnyKey,wouldLikeToViewGraphs=view,x=2,y=2,node_size=500)
    #                 # message4
    #                 message4 = "\nEnd node: " + current_node + "\nBFS found a path from start node " + node + " and end node " + current_node + "\nNew Edge: " + str(temp_list) + "\nTotal path: \n" + str(bfs_to_2d.flatten()) 
    #                 print(message4)
    #                 addToDoc(document,text=message4,addparagraph=True)
    #                 document.add_page_break()
    #                 temp_list = []
            
def task2(document):
    #current task
    task = "Task 1 "
    #temp file for graphs
    tmpFile = "graph.png"
    # move on message
    #pressAnyKey = "\n\n" + task + "graph\n\n" + "\n---PRESS ANY KEY TO MOVE ON!---\n\n"
    # questions
    question1 = "\n\nQUESTION 1:\nUse an application to find the strongly connected components of the digraph.\n\n"
    question2 = "\n\nQUESTION 2:\nDraw the digraph as a 'meta graph' of its strongly connected components in the report.\n\n"
    question3 = "\n\nQUESTION 3:\nRepresent the 'meta graph' as a DAG and linearize it in its topological order.\n\n"
    #create graph
    G=nx.DiGraph()
    
    G_directed_edges = [(1,3),
                       (3,2),
                       (3,5),
                       (2,1),
                       (4,1),
                       (4,2),
                       (4,12),
                       (5,6),
                       (5,8),
                       (6,8),
                       (6,7),
                       (8,9),
                       (9,5),
                       (9,11),
                       (8,10),
                       (10,9),
                       (6,10),
                       (7,10),
                       (10,11),
                       (11,12)]
    #add edges
    G.add_edges_from(G_directed_edges)
    # digraph 1

    # #get connected components
    # conn_comp = getConnectedComponents(G)

    pos={ # positions for nodes
        
        1:(0,10),
        3:(2.5,10),
        5:(5,10),
        6:(7.5,10),
        4:(0,7.5),
        2:(2.5,7.5),
        9:(5,7.5),
        8:(6.2,7.5),
        7:(9,7.5),
        12:(2.5,5),
        11:(5,5),
        10:(7.5,5),
    }

    #
    # ######################### --------------- QUESTION 1 ---------------#########################  #
    # 
    print(question1)
    task = "task 2 question 1"

    # Use an application to find the strongly connected components of the digraph
    #ask if the user wants to compare the graphs, or just view it later in the file
    view,answerQuestion = viewGraphs()

    if answerQuestion:
        # generate undirected graph 
        generateGraph(G,pos,tmpFile,document,task,mainIllustration=True,wouldLikeToViewGraphs=True)
        # add question to document
        addToDoc(document,text=question1,addparagraph=True)
        document.add_page_break()

        #get scc's from digraph
        scc = getStronglyConnectedComponents(G)
        #message1
        message1 = "Strongly Connected Components for the digraph: "
        for scc_group in scc:
            message1 += "\n" + str(scc_group)
        print(message1)
        addToDoc(document,text=message1,addparagraph=True)
        document.add_page_break()

    #
    # ######################### --------------- QUESTION 2 ---------------#########################  #
    # 
    print(question2)

    task = "task 2 question 2"

    # Draw the digraph as a 'meta graph' of its strongly connected components in the report
    
    #ask if the user wants to compare the graphs, or just view it later in the file
    view,answerQuestion = viewGraphs()

    if answerQuestion:
        # generate undirected graph 
        generateGraph(G,pos,tmpFile,document,task,mainIllustration=True,wouldLikeToViewGraphs=True)
        # add question to document
        addToDoc(document,text=question2,addparagraph=True)
        document.add_page_break()

        meta_graph = createDigraphWithSCC(G)
        #message1
        message1 = "Digraph as a meta-graph of its strongly connected components"
        print(message1)
        addToDoc(document,text=message1,addparagraph=True)
        answer = "Digraph: meta graph of its SCC"
        generateGraph(meta_graph,pos,tmpFile,document,task,answer=answer,mainIllustration=False,wouldLikeToViewGraphs=True)
        document.add_page_break()
    
    #
    # ######################### --------------- QUESTION 3 ---------------#########################  #
    # 
    print(question3)
    task = "task 2 question 3"
    # Represent the 'meta graph' as a DAG and linearize it in its topological order

    #ask if the user wants to compare the graphs, or just view it later in the file
    view,answerQuestion = viewGraphs()

    if answerQuestion:
        # generate undirected graph 
        generateGraph(G,pos,tmpFile,document,task,mainIllustration=True,wouldLikeToViewGraphs=True)
        # add question to document
        addToDoc(document,text=question3,addparagraph=True)
        document.add_page_break()

        meta_graph = createDigraphWithSCC(G)
        DAGgraph,pos = createDAG(meta_graph)
        #message1
        message1 = "Meta-graph as a DAG, linearized in its topological order"
        print(message1)
        addToDoc(document,text=message1,addparagraph=True)
        answer = "DAG: linearized in topological order"
        generateGraph(DAGgraph,pos,tmpFile,document,task,answer=answer,mainIllustration=False,wouldLikeToViewGraphs=True)
        document.add_page_break() 


def task3(document):
    #current task
    task = "Task 3"
    #temp file for graphs
    tmpFile = "graph.png"
    # move on message
    #pressAnyKey = "\n\n" + task + "graph\n\n" + "\n---PRESS ANY KEY TO MOVE ON!---\n\n"
    # questions
    question1 = "\n\nQUESTION 1:\nWrite an application that applies Dijkstra’s algorithm to produce the shortest path tree for" + \
                " a weighted graph with a given starting node. Test and verify your program with the given graph starting with node A\n\n"
    question2 = "\n\nQUESTION 2:\nWrite a program that produces a minimum spanning tree for a connected weighted graph.\n\n"
    question3 = "\n\nQUESTION 3:\nAre a shortest path tree and a minimum spanning tree usually the same?\n\n"
    question4 = "\n\nQUESTION 4:\nIf the graph has an edge with a negative weight, can you apply Dijkstra's algorithm to find a shortest path tree?\n\n"


    G_weighted_undirected = nx.Graph()

    G_weighted_undirected.add_edge("A","B",weight=22,color='black')
    G_weighted_undirected.add_edge("A","C",weight=9,color='black')
    G_weighted_undirected.add_edge("A","D",weight=12,color='black')
    G_weighted_undirected.add_edge("B","F",weight=36,color='black')
    G_weighted_undirected.add_edge("C","B",weight=35,color='black')
    G_weighted_undirected.add_edge("C","D",weight=4,color='black')
    G_weighted_undirected.add_edge("C","E",weight=65,color='black')
    G_weighted_undirected.add_edge("C","F",weight=42,color='black')
    G_weighted_undirected.add_edge("D","E",weight=33,color='black')
    G_weighted_undirected.add_edge("E","G",weight=23,color='black')
    G_weighted_undirected.add_edge("E","F",weight=18,color='black')
    G_weighted_undirected.add_edge("F","G",weight=39,color='black')
    G_weighted_undirected.add_edge("B","H",weight=34,color='black')
    G_weighted_undirected.add_edge("D","I",weight=30,color='black')
    G_weighted_undirected.add_edge("H","I",weight=19,color='black')
    G_weighted_undirected.add_edge("F","H",weight=24,color='black')
    G_weighted_undirected.add_edge("G","I",weight=21,color='black')
    G_weighted_undirected.add_edge("G","H",weight=25,color='black')


    pos={ # positions for nodes
        
        "A":(0,7.5),
        "B":(2.5,10),
        "C":(2.5,7.5),
        "D":(2.5,2.5),
        "E":(5,5),
        "F":(5,7.5),
        "G":(7.5,5),
        "H":(10,10),
        "I":(10,2.5),
    }

    #
    # ######################### --------------- QUESTION 1 ---------------#########################  #
    # 
    print(question1)
    task = "task 3 question 1"
    # Write an application that applies Dijkstra’s algorithm to produce the shortest path 
    #tree for a weighted graph with a given starting node. Test and verify your program 
    #with the given graph starting with node A
    #ask if the user wants to compare the graphs, or just view it later in the file
    view,answerQuestion = viewGraphs()     

    if answerQuestion:
        # generate undirected graph 
        generateGraph(G_weighted_undirected,pos,tmpFile,document,task,mainIllustration=True,wouldLikeToViewGraphs=True)
        # add question to document
        addToDoc(document,text=question1,addparagraph=True) 
        document.add_page_break()

        answer = "Dijkstra’s algorithm: spt graph"
        start_node = "A"
        generateGraphForSPT(document,G_weighted_undirected,pos,task,answer,tmpFile,start_node)
        
        
    #
    # ######################### --------------- QUESTION 2 ---------------#########################  #
    # 
    print(question2)
    task = "task 3 question 2"
    
    # Write a program that produces a minimum spanning tree for a connected weighted graph
    
    #ask if the user wants to compare the graphs, or just view it later in the file
    view,answerQuestion = viewGraphs()     

    if answerQuestion:
        # generate undirected graph 
        generateGraph(G_weighted_undirected,pos,tmpFile,document,task,mainIllustration=True,wouldLikeToViewGraphs=True)
        # add question to document
        addToDoc(document,text=question2,addparagraph=True)
        document.add_page_break()

        answer = "Kruskal's algorithm: mst graph"
        generateGraphForMST(document,G_weighted_undirected,pos,task,answer,tmpFile)
                    
        #REVISE: old method....
        # answer = "mst graph"
        # mst = nx.minimum_spanning_tree(G_weighted_undirected,weight='weight',algorithm='kruskal')
        # mst = nx.minimum
        # edges = mst.edges()
        # #message1 
        # message1 = "Minimum Spanning Tree using Kruskal's algorithm\n" + "Edges: \n" + str(edges) + "\nGraph: "
        # print(message1)
        # addToDoc(document,text=message1,addparagraph=True)
        # generateGraph(mst,pos,tmpFile,document,task,answer=answer,mainIllustration=False,wouldLikeToViewGraphs=True)
        # document.add_page_break()

    #
    # ######################### --------------- QUESTION 3 ---------------#########################  #
    # 
    print(question3)
    task = "task 3 question 3"

    # Are a shortest path tree and a minimum spanning tree usually the same?

    #ask if the user wants to compare the graphs, or just view it later in the file
    view,answerQuestion = viewGraphs()

    if answerQuestion:
        # generate undirected graph 
        generateGraph(G_weighted_undirected,pos,tmpFile,document,task,mainIllustration=True,wouldLikeToViewGraphs=True)
        # add question to document
        addToDoc(document,text=question3,addparagraph=True)
        document.add_page_break()

        operation = "compare spt and mst"
        answer = "Dijkstra’s algorithm: spt graph"
        start_node = "A"
        G_spt = generateGraphForSPT(document,G_weighted_undirected,pos,task,answer,tmpFile,start_node,viewGraph=False,affectDocument=False)
        answer = "Kruskal's algorithm: mst graph"
        G_mst = generateGraphForMST(document,G_weighted_undirected,pos,task,answer,tmpFile,viewGraph=False,affectDocument=False)
        answer = "spt vs mst"
        generateGraphSideBySideForComparison(document,G_spt,G_mst,pos,task,tmpFile,operation,answer=answer,wouldLikeToViewGraphs=True)
        message = "\nNo: \nSPT identifies a path to reach a destination using the shortest path (with a given start node)."
        message += "\nMST finds the least weighted path connecting each node with no destination" 
        addToDoc(document,text=message,addparagraph=True)
        document.add_page_break()

    #
    # ######################### --------------- QUESTION 4 ---------------#########################  #
    # 
    print(question4)
    task = "task 3 question 4"

    # If the graph has an edge with a negative weight, can you apply Dijkstra's algorithm to find a shortest path tree?

    #ask if the user wants to compare the graphs, or just view it later in the file
    view,answerQuestion = viewGraphs()
    
    if answerQuestion:
        #add negative edge
        G_error_prone = G_weighted_undirected.copy()
        G_error_prone.add_edge("G","I",weight=-21,color='black')
        # generate undirected graph 
        generateGraph(G_error_prone,pos,tmpFile,document,task,mainIllustration=True,wouldLikeToViewGraphs=True)
        # add question to document
        addToDoc(document,text=question4,addparagraph=True)
        document.add_page_break()

        note = "\nCurrent graph: negative edge [I,G]"
        print(note)
        generateGraph(G_error_prone,pos,tmpFile,document,task,answer=note,mainIllustration=True,wouldLikeToViewGraphs=True)
        start_node = "A"
        message = "\nStart node: " + start_node + "\nEnd node: NONE GIVEN"
        message +="\nIf the graph has an edge with a negative weight (ie., [I,G]), and provided with a start node and NO END NODE, and in that shortest path the negative edge is met," 
        message += " then Dijkstra's algorithm cannot be applied to find spt\nValueError:\n"
        try:
            spt = nx.single_source_dijkstra_path(G_error_prone,start_node,weight='weight')
            nx.draw(spt,pos=pos,node_size=700,with_labels=True,width=3)
        except ValueError as e:
            # print(e.with_traceback())
            message += str(e) #error message in the document.. without traceback
        print(message)
        addToDoc(document,text=message,addparagraph=True)
        document.add_page_break()
        
        note = "\nCurrent graph: negative edge [I,G]"
        print(note)
        generateGraph(G_error_prone,pos,tmpFile,document,task,answer=note,mainIllustration=True,wouldLikeToViewGraphs=True)
        message = "\nStart node: " + start_node + "\nTarget node: G"
        message +="\nIf the graph has an edge with a negative weight (ie., [I,G]), and provided with a start node and END NODE, and in that shortest path the negative edge is met AT THE DESTINATION" 
        message += " but the algorithm is told to stop at that destination, then Dijkstra's algorithm can be applied to find spt\n" 
        answer = "Dijkstra’s algorithm: spt graph"
        start_node = "A"
        target_node = "G"
        print(message)
        generateGraphForSPT(document,G_error_prone,pos,task,answer,tmpFile,start_node,target_node=target_node)
        addToDoc(document,text=message,addparagraph=True)
        document.add_page_break()

        note = "\nCurrent graph: negative edges [A,D],[I,G]"
        print(note)
        G_error_prone.add_edge("A","D",weight=-12,color='black')
        generateGraph(G_error_prone,pos,tmpFile,document,task,answer=note,mainIllustration=True,wouldLikeToViewGraphs=True)
        message = "\nStart node: " + start_node + "\nTarget node: G"
        message +="\nIf the graph has edges with a negative weight (ie., [A,D], [I,G]), and provided with a start node and END NODE, and in that shortest path the negative edge is met BEFORE THE DESTINATION," 
        message += " then Dijkstra's algorithm cannot be applied to find spt\nValueError:\n" 
        answer = "Dijkstra’s algorithm: spt graph"
        start_node = "A"
        target_node = "G"
        print(message)
        try:
            generateGraphForSPT(document,G_error_prone,pos,task,answer,tmpFile,start_node,target_node=target_node)
        except ValueError as e:
            message += str(e)
            print(message)
        addToDoc(document,text=message,addparagraph=True)
        document.add_page_break()

    
    
    

    


# main driver
def main():
    #get task 1 undirected graph
    document = Document()
    document.add_heading('Test Results on Graph Algorithms',level=0)
    document.add_page_break()


    choice = True
    showInvalidinput = False
    createResults = False
    while choice:
        #get input and ask which task to view
        print("\nPlease select the task to be viewed:\n1. DFS & BFS on undirected graph\n2. Connected digraph as a 'meta graph'" \
              + "\n3. Dijkstra and Kruskal's algorithms on a weighted undirected graph \n4. Quit/Save Results")
        opt1 = input("\nChoice: ")
        if opt1.__eq__("1"): 
            task1(document=document)
        elif opt1.__eq__("2"):
            task2(document=document)
        elif opt1.__eq__("3"):
            task3(document=document)
        elif opt1.__eq__("4"):
            print("Would you like to save these results?\n1. Yes\n2. No")
            opt2 = input("\nChoice: ")
            if opt2.__eq__("1"):
                choice = False
                createResults = True
            elif opt2.__eq__("2"):
                choice = False
            else:
                showInvalidinput = True
        else:
            showInvalidinput = True
        # show when invalid input
        if showInvalidinput:
            print("\n\nPlease select a valid option\n\n")
            choice = True
            showInvalidinput = False
            createResults = False

    if createResults:
        # # do stuff here
        cwd = os.getcwd()
        print("Creating document...")
        #document = Document(cwd + "Task Results.docx")
        filename = "Task Results.docx"
        try_again = True
        i = 1
        #in case the file is opened, to prevent loss of runs, create another file
        while try_again:
            try:
                document.save(filename)
                try_again = False
            except:
                filename = "Task Results(" + str(i) +").docx"
                i += 1
        # remove graph.png files
        for file in os.listdir("."):
            if file.endswith("graph.png") or file.endswith("mstgraph.png") or file.endswith("sideByside.png") or file.endswith("sptgraph.png"):
                os.remove(file)
        if i>1:
            print("\n\nYou had a similar file name opened, so another is going to be created with (" + str(i-1)+") in the filename.")
        print("\n\nDocument can be found in the directory:\n\n" + cwd + "\\" + filename)

    print("\n\nThank you! Take care.")


if __name__ == '__main__':
    main()